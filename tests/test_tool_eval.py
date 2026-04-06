"""Tests for multi-tool Markdown evaluation."""

from pathlib import Path

from docx import Document

from parserx.tool_eval.adapters import BuiltinDocPdfAdapter
from parserx.tool_eval.runner import MultiToolEvalRunner


def test_builtin_doc_pdf_adapter_converts_docx_with_headings_and_tables(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("This is a short summary.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Warnings"
    table.rows[1].cells[1].text = "2"
    document.save(docx_path)

    result = BuiltinDocPdfAdapter().parse(docx_path, tmp_path / "artifacts")

    assert "# Project Overview" in result.markdown
    assert "This is a short summary." in result.markdown
    assert "| Metric | Value |" in result.markdown
    assert result.metadata["parser"] == "python-docx"


def test_multi_tool_runner_artifact_only_without_expected_md(tmp_path: Path):
    """tool-eval should produce artifacts even when expected.md is missing."""
    # Create a doc dir with an input file but no expected.md
    doc_dir = tmp_path / "gt" / "no_gt_doc"
    doc_dir.mkdir(parents=True)
    docx_path = doc_dir / "input.docx"
    document = Document()
    document.add_heading("Title", level=1)
    document.add_paragraph("Body text.")
    document.save(docx_path)

    artifacts_dir = tmp_path / "artifacts"
    runner = MultiToolEvalRunner(tools=[BuiltinDocPdfAdapter()])
    records = runner.evaluate_dir(tmp_path / "gt", artifacts_dir)

    assert len(records) == 1
    record = records[0]
    assert record.status == "artifact_only"
    assert record.metrics is None
    assert Path(record.output_path).exists()
    output_md = Path(record.output_path).read_text(encoding="utf-8")
    assert "Title" in output_md

    # Report should render without errors
    report = MultiToolEvalRunner.format_report(
        records,
        ground_truth_dir=tmp_path / "gt",
        artifacts_dir=artifacts_dir,
    )
    assert "artifact_only" in report


def test_multi_tool_runner_writes_artifacts_and_manifest(tmp_path: Path):
    ground_truth_dir = Path("ground_truth_public/basic_report")
    artifacts_dir = tmp_path / "tool_eval_artifacts"
    runner = MultiToolEvalRunner(tools=[BuiltinDocPdfAdapter()])

    records = runner.evaluate_dir(ground_truth_dir, artifacts_dir)

    assert len(records) == 1
    record = records[0]
    assert record.status == "ok"
    assert Path(record.output_path).exists()
    assert (artifacts_dir / "builtin_doc_pdf" / "basic_report" / "metadata.json").exists()
    assert (artifacts_dir / "manifest.json").exists()

    report = MultiToolEvalRunner.format_report(
        records,
        ground_truth_dir=ground_truth_dir,
        artifacts_dir=artifacts_dir,
    )
    assert "Multi-Tool Markdown Evaluation" in report
    assert "builtin_doc_pdf" in report
    assert "basic_report" in report
